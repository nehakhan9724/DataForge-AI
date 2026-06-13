import io
import json
import mimetypes
import os
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

load_dotenv(Path(__file__).with_name('.env'))

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from google import genai
except ImportError:
    genai = None

try:
    from PIL import Image, UnidentifiedImageError
except ImportError:
    Image = None
    UnidentifiedImageError = OSError

APP_TITLE = 'Visual Data Assistant API'
APP_VERSION = '1.0.0'
DEFAULT_MODEL = os.getenv('GEMINI_MODEL', 'gemini-2.5-flash').strip() or 'gemini-2.5-flash'
MAX_UPLOAD_BYTES = 200 * 1024 * 1024
MAX_MODEL_TEXT_CHARS = 120000
MAX_OUTPUT_TOKENS = 8192
MAX_FIELDS = 48
MAX_TABLES = 4
MAX_TABLE_ROWS = 80
MAX_SPREADSHEET_MODEL_ROWS = 220
MAX_SPREADSHEET_MODEL_COLS = 40
ALLOWED_MODES = {'smart', 'key_value', 'table'}

SUPPORTED_EXTENSIONS = {
    '.pdf': 'pdf',
    '.png': 'image',
    '.jpg': 'image',
    '.jpeg': 'image',
    '.webp': 'image',
    '.bmp': 'image',
    '.gif': 'image',
    '.txt': 'text',
    '.md': 'text',
    '.xml': 'xml',
    '.xlsx': 'spreadsheet',
    '.xls': 'spreadsheet',
    '.docx': 'docx',
}

OUTPUT_SCHEMA = '{"status":"success|unreadable|no_data","message":"string","title":"string","summary_text":"string","warnings":["string"],"key_value_pairs":[{"field_label":"string","value":"string","normalized_value":"string","value_type":"text|number|date|currency|percentage|boolean|unknown","confidence":"high|medium|low"}],"tables":[{"table_name":"string","columns":["string"],"rows":[{"column":"value"}]}]}'

BOOLEAN_WORDS = {'yes', 'no', 'true', 'false', 'present', 'absent', 'pass', 'fail'}


class TextExtractRequest(BaseModel):
    text: str
    extraction_mode: str = 'smart'
    title: str = 'Pasted Text'


app = FastAPI(title=APP_TITLE, version=APP_VERSION)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[origin.strip() for origin in os.getenv('CORS_ORIGINS', '*').split(',') if origin.strip()] or ['*'],
    allow_credentials=True,
    allow_methods=['*'],
    allow_headers=['*'],
)


def get_client():
    if genai is None:
        raise HTTPException(status_code=500, detail='Missing dependency: google-genai is not installed.')
    api_key = os.getenv('GEMINI_API_KEY', '').strip()
    if not api_key:
        raise HTTPException(status_code=500, detail='Missing GEMINI_API_KEY in backend environment.')
    return genai.Client(api_key=api_key)


def clean_text(value: Any) -> str:
    return '' if value is None else str(value).strip()


def to_snake_case(value: str) -> str:
    value = re.sub(r'[^a-zA-Z0-9]+', '_', clean_text(value).lower())
    value = re.sub(r'_+', '_', value).strip('_')
    return value or 'field'


def human_size(num_bytes: int) -> str:
    size = float(num_bytes)
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size < 1024 or unit == 'GB':
            return f'{size:.1f} {unit}'
        size /= 1024
    return f'{num_bytes} B'


def strip_code_fences(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith('```'):
        cleaned = re.sub(r'^```[a-zA-Z0-9_-]*\n?', '', cleaned)
    if cleaned.endswith('```'):
        cleaned = cleaned[:-3]
    return cleaned.strip()


def extract_balanced_json_object(text: str) -> str:
    start = text.find('{')
    if start == -1:
        return ''
    depth = 0
    in_string = False
    escaped = False
    for index, char in enumerate(text[start:], start=start):
        if in_string:
            if escaped:
                escaped = False
            elif char == '\\':
                escaped = True
            elif char == '"':
                in_string = False
            continue
        if char == '"':
            in_string = True
        elif char == '{':
            depth += 1
        elif char == '}':
            depth -= 1
            if depth == 0:
                return text[start : index + 1]
    return ''


def normalize_json_candidate(text: str) -> str:
    candidate = strip_code_fences(text)
    candidate = candidate.strip()
    candidate = re.sub(r',\s*([}\]])', r'\1', candidate)
    return candidate


def parse_model_json(text: str) -> Dict[str, Any]:
    cleaned = strip_code_fences(text)
    if not cleaned:
        raise HTTPException(status_code=502, detail='Model returned an empty response.')

    candidates: List[str] = []
    normalized = normalize_json_candidate(cleaned)
    if normalized:
        candidates.append(normalized)

    balanced = extract_balanced_json_object(cleaned)
    if balanced:
        balanced = normalize_json_candidate(balanced)
        if balanced and balanced not in candidates:
            candidates.append(balanced)

    for candidate in candidates:
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            continue

    try:
        json.loads(normalized)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail=f'Model did not return valid JSON: {exc}') from exc

    raise HTTPException(status_code=502, detail='Model did not return valid JSON.')


def strict_retry_prompt(prompt: str) -> str:
    return (
        f"{prompt} "
        'IMPORTANT: Return STRICT valid minified JSON only. No markdown, no prose, no comments, no trailing commas. ' 
        'Every opened quote, bracket, and brace must be closed. ' 
        'If the result is large, keep only the most relevant fields and table rows, but the JSON must remain complete and valid.'
    )


def generate_json_response(*, client: Any, model: str, contents: List[Any]) -> Dict[str, Any]:
    last_error = None
    for attempt in range(2):
        response = client.models.generate_content(
            model=model,
            contents=contents,
            config={'response_mime_type': 'application/json', 'temperature': 0, 'max_output_tokens': MAX_OUTPUT_TOKENS},
        )
        response_text = getattr(response, 'text', '')
        try:
            return parse_model_json(response_text)
        except HTTPException as exc:
            last_error = exc.detail
            if attempt == 1:
                break
    raise HTTPException(status_code=502, detail=f'Model request failed: {last_error or "Invalid model response."}')


def guess_extension(filename: str, content_type: Optional[str]) -> str:
    ext = Path(filename or '').suffix.lower()
    return ext or (mimetypes.guess_extension(content_type or '') or '').lower()


def detect_kind(filename: str, content_type: Optional[str]) -> Tuple[str, str]:
    ext = guess_extension(filename, content_type)
    kind = SUPPORTED_EXTENSIONS.get(ext)
    if not kind:
        raise HTTPException(status_code=400, detail='Unsupported file type. Use PDF, image, TXT, MD, XML, XLS/XLSX, or DOCX.')
    return ext, content_type or mimetypes.guess_type(filename)[0] or 'application/octet-stream'


def ensure_image_is_valid(file_bytes: bytes, filename: str) -> None:
    if Image is None:
        return
    try:
        with Image.open(io.BytesIO(file_bytes)) as image:
            image.verify()
    except (UnidentifiedImageError, OSError) as exc:
        raise HTTPException(status_code=400, detail=f'{filename}: Please upload a valid image or clearer file.') from exc


def preprocess_image_for_model(file_bytes: bytes, filename: str) -> bytes:
    if Image is None:
        return file_bytes
    suffix = Path(filename).suffix.lower()
    if suffix not in {'.png', '.jpg', '.jpeg', '.webp', '.bmp'}:
        return file_bytes
    try:
        with Image.open(io.BytesIO(file_bytes)) as image:
            image = image.convert('RGB')
            width, height = image.size
            max_edge = 1700
            scale = min(1.0, max_edge / max(width, height))
            if scale < 1.0:
                image = image.resize((max(1, int(width * scale)), max(1, int(height * scale))))
            buffer = io.BytesIO()
            image.save(buffer, format='JPEG', quality=84, optimize=True)
            optimized = buffer.getvalue()
            return optimized if len(optimized) < len(file_bytes) else file_bytes
    except Exception:
        return file_bytes


def build_prompt(filename: str, file_kind: str, extraction_mode: str) -> str:
    mode_hint = {
        'key_value': 'Prioritize exact labeled fields, marks, totals, ids, dates, and short values.',
        'table': 'Prioritize exact row and column reconstruction when the content is tabular.',
    }.get(extraction_mode, 'Return both fields and tables when they are explicitly present.')
    return (
        f'Extract structured data from the attached {file_kind} named "{filename}". '
        f'{mode_hint} '
        'Do not invent values, infer missing values, or paraphrase cell contents. '
        'If the file is unclear, use status="unreadable" and a short message asking for a clearer file. '
        'If it contains no usable structured data, use status="no_data" and message="This file doesn\'t contain any data." '
        'If a pair is reversed like "268 : bananas", normalize it to label="bananas", value="268" when obvious. '
        f'Return JSON only in this exact shape: {OUTPUT_SCHEMA} '
        f'Keep it concise but complete: at most {MAX_FIELDS} key_value_pairs, {MAX_TABLES} tables, and {MAX_TABLE_ROWS} rows per table. '
        'Return strict valid minified JSON only. No markdown. No commentary. Leave warnings empty unless absolutely necessary.'
    )


def prepare_text_for_model(text: str) -> str:
    text = clean_text(text).replace('\x00', '')
    lines = [re.sub(r'\s+', ' ', line).strip() for line in text.replace('\r', '\n').split('\n')]
    compact = '\n'.join(line for line in lines if line)
    if len(compact) <= MAX_MODEL_TEXT_CHARS:
        return compact

    window = MAX_MODEL_TEXT_CHARS // 3
    head = compact[:window]
    middle_start = max(window, len(compact) // 2 - window // 2)
    middle = compact[middle_start : middle_start + window]
    tail = compact[-window:]
    return f'{head}\n...\n{middle}\n...\n{tail}'


def normalize_scalar(value: Any) -> str:
    if value is None:
        return ''
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f'{value:.2f}'.rstrip('0').rstrip('.')
    if isinstance(value, (int, bool)):
        return str(value)
    return clean_text(value)


def infer_value_type(value: str) -> str:
    stripped = clean_text(value)
    if not stripped:
        return 'unknown'
    if re.fullmatch(r'\d{4}[-/]\d{1,2}[-/]\d{1,2}', stripped):
        return 'date'
    if re.fullmatch(r'[-+]?\d+(\.\d+)?%', stripped):
        return 'percentage'
    if re.fullmatch(r'[$€£₹]\s*[-+]?\d[\d,]*(\.\d+)?', stripped) or re.fullmatch(r'[-+]?\d[\d,]*(\.\d+)?\s*(usd|inr|eur|gbp)', stripped.lower()):
        return 'currency'
    if re.fullmatch(r'[-+]?\d[\d,]*(\.\d+)?', stripped):
        return 'number'
    if stripped.lower() in BOOLEAN_WORDS:
        return 'boolean'
    return 'text'


def build_pair(label: str, value: str, confidence: str = 'medium') -> Dict[str, Any]:
    normalized = clean_text(value)
    return {
        'field_label': clean_text(label) or 'Unnamed Field',
        'field_key': to_snake_case(label),
        'value': clean_text(value),
        'normalized_value': normalized,
        'value_type': infer_value_type(normalized or value),
        'confidence': confidence if confidence in {'high', 'medium', 'low'} else 'medium',
        'source_text': f'{clean_text(label)}: {clean_text(value)}',
    }


def dedupe_pairs(pairs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = set()
    output = []
    for pair in pairs:
        label = clean_text(pair.get('field_label'))
        value = clean_text(pair.get('normalized_value') or pair.get('value'))
        key = (label.lower(), value.lower())
        if label and value and key not in seen:
            seen.add(key)
            output.append(pair)
    return output[:MAX_FIELDS]


def cleanup_dataframe(df):
    if pd is None:
        raise HTTPException(status_code=500, detail='Missing dependency: pandas is not installed.')
    if df is None or df.empty:
        return pd.DataFrame()
    df = df.copy().dropna(axis=0, how='all').dropna(axis=1, how='all')
    if df.empty:
        return df
    df.columns = [clean_text(c) or f'column_{i + 1}' for i, c in enumerate(df.columns)]
    return df.reset_index(drop=True)


def df_to_table_payload(df, table_name: str) -> Dict[str, Any]:
    df = cleanup_dataframe(df)
    if df.empty:
        return {'table_name': table_name, 'columns': [], 'rows': []}
    if len(df) > MAX_TABLE_ROWS:
        df = df.head(MAX_TABLE_ROWS)
    columns = [clean_text(col) or f'column_{idx + 1}' for idx, col in enumerate(df.columns)]
    rows = []
    for _, row in df.iterrows():
        rows.append({column: normalize_scalar(row.get(column)) for column in columns})
    return {'table_name': table_name, 'columns': columns, 'rows': rows}


def infer_pairs_from_dataframe(df) -> List[Dict[str, Any]]:
    df = cleanup_dataframe(df)
    pairs: List[Dict[str, Any]] = []
    if df.empty:
        return pairs
    if len(df.columns) == 2:
        left, right = list(df.columns)
        for _, row in df.iterrows():
            label = clean_text(row.get(left))
            value = normalize_scalar(row.get(right))
            if label and value:
                pairs.append(build_pair(label, value, 'high'))
        return dedupe_pairs(pairs)
    if len(df) == 1 and len(df.columns) <= 24:
        row = df.iloc[0]
        for column in df.columns:
            value = normalize_scalar(row.get(column))
            if value:
                pairs.append(build_pair(column, value, 'high'))
    return dedupe_pairs(pairs)


def parse_spreadsheet(file_bytes: bytes, ext: str) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    if pd is None:
        raise HTTPException(status_code=500, detail='Missing dependency: pandas is not installed.')
    engine = 'openpyxl' if ext == '.xlsx' else 'xlrd'
    sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, engine=engine)
    tables: List[Dict[str, Any]] = []
    pairs: List[Dict[str, Any]] = []
    for name, frame in sheets.items():
        clean_df = cleanup_dataframe(frame)
        if clean_df.empty:
            continue
        tables.append(df_to_table_payload(clean_df, clean_text(name) or 'Sheet'))
        pairs.extend(infer_pairs_from_dataframe(clean_df))
    return tables[:MAX_TABLES], dedupe_pairs(pairs)


def spreadsheet_to_model_text(file_bytes: bytes, ext: str) -> Tuple[str, List[Dict[str, Any]], List[Dict[str, Any]]]:
    if pd is None:
        raise HTTPException(status_code=500, detail='Missing dependency: pandas is not installed.')
    engine = 'openpyxl' if ext == '.xlsx' else 'xlrd'
    sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, engine=engine)
    sections: List[str] = []
    fallback_tables: List[Dict[str, Any]] = []
    fallback_pairs: List[Dict[str, Any]] = []

    for name, frame in sheets.items():
        clean_df = cleanup_dataframe(frame)
        if clean_df.empty:
            continue

        fallback_tables.append(df_to_table_payload(clean_df, clean_text(name) or 'Sheet'))
        fallback_pairs.extend(infer_pairs_from_dataframe(clean_df))

        limited_df = clean_df.iloc[:MAX_SPREADSHEET_MODEL_ROWS, :MAX_SPREADSHEET_MODEL_COLS].copy()
        columns = [clean_text(col) or f'column_{idx + 1}' for idx, col in enumerate(limited_df.columns)]

        sections.append(f'[Sheet] {clean_text(name) or "Sheet"}')
        sections.append(f'Row count: {len(clean_df)} | Column count: {len(clean_df.columns)}')
        sections.append('Columns:	' + '	'.join(columns))

        for _, row in limited_df.iterrows():
            sections.append('	'.join(normalize_scalar(row.get(column)) for column in columns))

        if len(clean_df) > len(limited_df):
            sections.append(f'... {len(clean_df) - len(limited_df)} more rows not shown')
        if len(clean_df.columns) > len(columns):
            sections.append(f'... {len(clean_df.columns) - len(columns)} more columns not shown')
        sections.append('')

    return prepare_text_for_model('\n'.join(sections)), fallback_tables[:MAX_TABLES], dedupe_pairs(fallback_pairs)


def extract_docx_text(file_bytes: bytes) -> str:
    if Document is None:
        raise HTTPException(status_code=500, detail='Missing dependency: python-docx is not installed.')
    document = Document(io.BytesIO(file_bytes))
    lines: List[str] = []
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if any(cells):
                lines.append('\t'.join(cells))
    return '\n'.join(lines).strip()


def decode_text_bytes(file_bytes: bytes) -> str:
    for encoding in ['utf-8', 'utf-16', 'latin-1']:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode('utf-8', errors='ignore')


def call_gemini_with_text(*, client: Any, model: str, prompt: str, text_content: str) -> Dict[str, Any]:
    attempts = [prompt, strict_retry_prompt(prompt)]
    last_error = None
    for attempt_prompt in attempts:
        try:
            return generate_json_response(client=client, model=model, contents=[attempt_prompt, text_content])
        except HTTPException as exc:
            last_error = exc
            continue
    raise last_error or HTTPException(status_code=502, detail='Model request failed.')


def call_gemini_with_uploaded_file(*, client: Any, model: str, prompt: str, file_bytes: bytes, filename: str) -> Dict[str, Any]:
    uploaded_file = None
    temp_path = None
    try:
        suffix = Path(filename or 'uploaded').suffix or '.bin'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name
        uploaded_file = client.files.upload(file=temp_path)
        attempts = [prompt, strict_retry_prompt(prompt)]
        last_error = None
        for attempt_prompt in attempts:
            try:
                return generate_json_response(client=client, model=model, contents=[attempt_prompt, uploaded_file])
            except HTTPException as exc:
                last_error = exc
                continue
        raise last_error or HTTPException(status_code=502, detail='Model request failed.')
    finally:
        if uploaded_file is not None:
            try:
                client.files.delete(name=uploaded_file.name)
            except Exception:
                pass
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except OSError:
                pass


def normalize_pair(pair: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    label = clean_text(pair.get('field_label'))
    value = clean_text(pair.get('normalized_value') or pair.get('value'))
    if not label and not value:
        return None
    confidence = clean_text(pair.get('confidence')).lower() or 'medium'
    if confidence not in {'high', 'medium', 'low'}:
        confidence = 'medium'
    value_type = clean_text(pair.get('value_type')).lower() or infer_value_type(value)
    if value_type not in {'text', 'number', 'date', 'currency', 'percentage', 'boolean', 'unknown'}:
        value_type = infer_value_type(value)
    return {
        'field_label': label or 'Unnamed Field',
        'field_key': to_snake_case(label or 'field'),
        'value': clean_text(pair.get('value')) or value,
        'normalized_value': value,
        'value_type': value_type,
        'confidence': confidence,
        'source_text': clean_text(pair.get('source_text')) or f'{label}: {value}',
    }


def normalize_table(table: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    columns = [clean_text(column) for column in table.get('columns', []) if clean_text(column)]
    rows: List[Dict[str, Any]] = []
    for row in table.get('rows', [])[:MAX_TABLE_ROWS]:
        if not isinstance(row, dict):
            continue
        normalized_row = {clean_text(k): clean_text(v) for k, v in row.items() if clean_text(k)}
        if not columns:
            columns = list(normalized_row.keys())
        if any(normalized_row.values()):
            rows.append({column: normalized_row.get(column, '') for column in columns})
    if not columns and not rows:
        return None
    return {
        'table_name': clean_text(table.get('table_name')) or 'Extracted Table',
        'columns': columns,
        'rows': rows[:MAX_TABLE_ROWS],
    }


def normalize_model_output(model_output: Dict[str, Any]) -> Dict[str, Any]:
    status = clean_text(model_output.get('status')).lower()
    if status not in {'success', 'unreadable', 'no_data'}:
        status = 'success'
    pairs = [item for raw in model_output.get('key_value_pairs', []) if isinstance(raw, dict) for item in [normalize_pair(raw)] if item]
    tables = [item for raw in model_output.get('tables', []) if isinstance(raw, dict) for item in [normalize_table(raw)] if item]
    pairs = dedupe_pairs(pairs)
    tables = tables[:MAX_TABLES]
    warnings = [clean_text(item) for item in model_output.get('warnings', []) if clean_text(item)][:3]
    if status == 'success' and not pairs and not tables:
        status = 'no_data'
    message = clean_text(model_output.get('message'))
    if status == 'unreadable' and not message:
        message = 'Please upload a valid image or clearer file.'
    if status == 'no_data' and not message:
        message = "This file doesn't contain any data."
    if status == 'success' and not message:
        message = 'Extraction completed successfully.'
    summary_text = clean_text(model_output.get('summary_text'))
    if status == 'success' and not summary_text:
        summary_text = f'Extracted {len(pairs)} field(s) and {len(tables)} table(s).'
    return {
        'status': status,
        'message': message,
        'title': clean_text(model_output.get('title')),
        'summary_text': summary_text,
        'warnings': warnings,
        'key_value_pairs': pairs,
        'tables': tables,
    }


def apply_mode_filters(normalized: Dict[str, Any], extraction_mode: str) -> Dict[str, Any]:
    output = dict(normalized)
    pairs = list(output.get('key_value_pairs', []))
    tables = list(output.get('tables', []))
    if extraction_mode == 'key_value':
        output['tables'] = []
        return output
    if extraction_mode == 'table':
        if not tables and pairs:
            tables = [{
                'table_name': 'Extracted Fields',
                'columns': ['Field', 'Value'],
                'rows': [{'Field': pair.get('field_label', ''), 'Value': pair.get('normalized_value') or pair.get('value') or ''} for pair in pairs[:MAX_TABLE_ROWS]],
            }]
        output['key_value_pairs'] = []
        output['tables'] = tables[:MAX_TABLES]
        return output
    return output


def build_summary_cards(status: str, key_value_pairs: List[Dict[str, Any]], tables: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    cards = [
        {'label': 'Status', 'value': status.replace('_', ' ').title()},
        {'label': 'Fields', 'value': str(len(key_value_pairs))},
        {'label': 'Tables', 'value': str(len(tables))},
    ]
    for pair in key_value_pairs[:3]:
        cards.append({'label': pair['field_label'], 'value': pair.get('normalized_value') or pair.get('value') or '-'})
    return cards[:6]


def build_final_response(*, filename: str, file_kind: str, file_bytes: bytes, extraction_mode: str, model: str, normalized: Dict[str, Any]) -> Dict[str, Any]:
    normalized = apply_mode_filters(normalized, extraction_mode)
    return {
        'message': normalized.get('message') or 'Done.',
        'result': {
            'status': normalized['status'],
            'message': normalized.get('message') or '',
            'title': normalized.get('title') or filename,
            'summary_text': normalized.get('summary_text') or '',
            'summary_cards': build_summary_cards(normalized['status'], normalized.get('key_value_pairs', []), normalized.get('tables', [])),
            'warnings': normalized.get('warnings', []),
            'key_value_pairs': normalized.get('key_value_pairs', []),
            'tables': normalized.get('tables', []),
            'meta': {
                'file_name': filename,
                'file_kind': file_kind,
                'file_size': human_size(len(file_bytes)),
                'model': model,
                'extraction_mode': extraction_mode,
            },
        },
    }


def build_spreadsheet_response(filename: str, ext: str, file_bytes: bytes, extraction_mode: str) -> Dict[str, Any]:
    spreadsheet_text, fallback_tables, fallback_pairs = spreadsheet_to_model_text(file_bytes, ext)
    if not spreadsheet_text and not fallback_tables and not fallback_pairs:
        normalized = {
            'status': 'no_data',
            'message': "This file doesn't contain any data.",
            'title': filename,
            'summary_text': 'No extractable rows or fields were found.',
            'warnings': [],
            'key_value_pairs': [],
            'tables': [],
        }
        return build_final_response(filename=filename, file_kind='spreadsheet', file_bytes=file_bytes, extraction_mode=extraction_mode, model=DEFAULT_MODEL, normalized=normalized)

    client = get_client()
    prompt = build_prompt(filename, 'spreadsheet', extraction_mode)

    try:
        normalized = normalize_model_output(call_gemini_with_text(client=client, model=DEFAULT_MODEL, prompt=prompt, text_content=spreadsheet_text))
        if normalized.get('status') == 'no_data' and (fallback_tables or fallback_pairs):
            raise HTTPException(status_code=502, detail='Model returned no usable spreadsheet data.')
        return build_final_response(filename=filename, file_kind='spreadsheet', file_bytes=file_bytes, extraction_mode=extraction_mode, model=DEFAULT_MODEL, normalized=normalized)
    except HTTPException:
        if fallback_tables or fallback_pairs:
            fallback_normalized = {
                'status': 'success',
                'message': 'Extraction completed successfully.',
                'title': filename,
                'summary_text': f'Extracted {len(fallback_pairs)} field(s) and {len(fallback_tables)} table(s).',
                'warnings': [],
                'key_value_pairs': fallback_pairs,
                'tables': fallback_tables,
            }
            return build_final_response(filename=filename, file_kind='spreadsheet', file_bytes=file_bytes, extraction_mode=extraction_mode, model=DEFAULT_MODEL, normalized=fallback_normalized)
        raise


def process_text_request(*, title: str, extraction_mode: str, text: str) -> Dict[str, Any]:
    prepared = prepare_text_for_model(text)
    if not prepared:
        normalized = {
            'status': 'no_data',
            'message': "This file doesn't contain any data.",
            'title': title,
            'summary_text': 'No extractable content was found.',
            'warnings': [],
            'key_value_pairs': [],
            'tables': [],
        }
        return build_final_response(filename=title, file_kind='text', file_bytes=b'', extraction_mode=extraction_mode, model=DEFAULT_MODEL, normalized=normalized)

    client = get_client()
    prompt = build_prompt(title, 'text', extraction_mode)
    normalized = normalize_model_output(call_gemini_with_text(client=client, model=DEFAULT_MODEL, prompt=prompt, text_content=prepared))
    return build_final_response(filename=title, file_kind='text', file_bytes=text.encode('utf-8', errors='ignore'), extraction_mode=extraction_mode, model=DEFAULT_MODEL, normalized=normalized)


@app.get('/health')
async def health() -> Dict[str, str]:
    return {'status': 'ok', 'app': APP_TITLE, 'version': APP_VERSION, 'model': DEFAULT_MODEL}


@app.post('/extract-text')
async def extract_text(request: TextExtractRequest) -> Dict[str, Any]:
    extraction_mode = clean_text(request.extraction_mode).lower() or 'smart'
    if extraction_mode not in ALLOWED_MODES:
        raise HTTPException(status_code=400, detail='extraction_mode must be smart, key_value, or table.')
    return process_text_request(title=clean_text(request.title) or 'Pasted Text', extraction_mode=extraction_mode, text=request.text)


@app.post('/extract')
async def extract_file(file: UploadFile = File(...), extraction_mode: str = Form('smart')) -> Dict[str, Any]:
    extraction_mode = clean_text(extraction_mode).lower() or 'smart'
    if extraction_mode not in ALLOWED_MODES:
        raise HTTPException(status_code=400, detail='extraction_mode must be smart, key_value, or table.')

    filename = file.filename or 'uploaded-file'
    ext, _content_type = detect_kind(filename, file.content_type)
    file_kind = SUPPORTED_EXTENSIONS[ext]
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail='Uploaded file is empty.')
    if len(file_bytes) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail='File is larger than the 200 MB app upload limit.')

    if file_kind == 'spreadsheet':
        return build_spreadsheet_response(filename, ext, file_bytes, extraction_mode)

    client = get_client()

    if file_kind == 'image':
        ensure_image_is_valid(file_bytes, filename)
        model_bytes = preprocess_image_for_model(file_bytes, filename)
        normalized = normalize_model_output(
            call_gemini_with_uploaded_file(client=client, model=DEFAULT_MODEL, prompt=build_prompt(filename, file_kind, extraction_mode), file_bytes=model_bytes, filename=filename)
        )
        return build_final_response(filename=filename, file_kind=file_kind, file_bytes=file_bytes, extraction_mode=extraction_mode, model=DEFAULT_MODEL, normalized=normalized)

    if file_kind == 'pdf':
        normalized = normalize_model_output(
            call_gemini_with_uploaded_file(client=client, model=DEFAULT_MODEL, prompt=build_prompt(filename, file_kind, extraction_mode), file_bytes=file_bytes, filename=filename)
        )
        return build_final_response(filename=filename, file_kind=file_kind, file_bytes=file_bytes, extraction_mode=extraction_mode, model=DEFAULT_MODEL, normalized=normalized)

    if file_kind == 'docx':
        text_content = extract_docx_text(file_bytes)
    else:
        text_content = decode_text_bytes(file_bytes)

    prepared_text = prepare_text_for_model(text_content)
    if not prepared_text:
        normalized = {
            'status': 'no_data',
            'message': "This file doesn't contain any data.",
            'title': filename,
            'summary_text': 'No extractable content was found.',
            'warnings': [],
            'key_value_pairs': [],
            'tables': [],
        }
        return build_final_response(filename=filename, file_kind=file_kind, file_bytes=file_bytes, extraction_mode=extraction_mode, model=DEFAULT_MODEL, normalized=normalized)

    normalized = normalize_model_output(
        call_gemini_with_text(client=client, model=DEFAULT_MODEL, prompt=build_prompt(filename, file_kind, extraction_mode), text_content=prepared_text)
    )
    return build_final_response(filename=filename, file_kind=file_kind, file_bytes=file_bytes, extraction_mode=extraction_mode, model=DEFAULT_MODEL, normalized=normalized)
